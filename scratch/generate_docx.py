import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level, space_before=12, space_after=6):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        run.font.name = 'Segoe UI'
        if level == 1:
            run.font.color.rgb = RGBColor(31, 78, 121) # Deep Blue
            run.font.size = Pt(18)
        elif level == 2:
            run.font.color.rgb = RGBColor(64, 64, 64)  # Charcoal
            run.font.size = Pt(14)
        elif level == 3:
            run.font.color.rgb = RGBColor(100, 100, 100) # Muted Gray
            run.font.size = Pt(12)
    return p

def create_document():
    doc = docx.Document()
    
    # Page Setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # ── TITLE PAGE ──────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(120)
    title_p.paragraph_format.space_after = Pt(10)
    run_title = title_p.add_run("EDUPORTAL AI")
    run_title.font.name = 'Segoe UI'
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(31, 78, 121)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(40)
    run_sub = subtitle_p.add_run("Desktop Lockdown Browser & Secure Exam Portal\nTechnical Implementation Plan & Architecture Specification")
    run_sub.font.name = 'Segoe UI'
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 100, 100)
    
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_before = Pt(150)
    run_meta = meta_p.add_run("Version: 1.6 (Decoupled Login & Session Launch)\nPrepared for: Engineering Management & Project Stakeholders\nStatus: Final Draft for Review\nTarget Platforms: Windows 10/11 & macOS 13+")
    run_meta.font.name = 'Segoe UI'
    run_meta.font.size = Pt(10.5)
    run_meta.font.color.rgb = RGBColor(64, 64, 64)
    
    doc.add_page_break()
    
    # ── 1. PROJECT EXECUTIVE SUMMARY ───────────────────────────────────────
    add_styled_heading(doc, "1. Executive Summary", level=1, space_before=24)
    
    p1 = doc.add_paragraph("EduPortal AI is an AI-powered student evaluation system utilizing Google Gemini (gemini-2.5-flash) and MongoDB. While the current prototype successfully supports multi-modal answer submissions (text, image, and voice), it operates in an open web environment without administrative proctoring controls or exam integrity limits.")
    p1.paragraph_format.space_after = Pt(8)
    p1.paragraph_format.line_spacing = 1.15
    
    p2 = doc.add_paragraph("This technical implementation plan details a two-tier strategy to harden the existing setup into a robust, secure exam portal:")
    p2.paragraph_format.space_after = Pt(6)
    
    bp1 = doc.add_paragraph(style='List Bullet')
    run_bp1_bold = bp1.add_run("Web Portal Upgrades: ")
    run_bp1_bold.bold = True
    bp1.add_run("Restructuring session routing, database models, scoped exam views, webcam integration, and proctoring metrics dashboards.")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    run_bp2_bold = bp2.add_run("PyQt6 Desktop Lockdown App: ")
    run_bp2_bold.bold = True
    bp2.add_run("A standalone desktop container that gates web portal access, suppresses operating system shortcuts, prevents display duplication, monitors focus states, and logs violation heartbeats.")
    
    p3 = doc.add_paragraph("By deploying this architecture, the platform will enforce that exams are only accessible within the secure container, ensuring complete visual and session integrity while tracking proctoring telemetry in real time.")
    p3.paragraph_format.space_before = Pt(10)
    p3.paragraph_format.space_after = Pt(18)
    
    # ── 2. TECHNICAL MODERNIZATION STRATEGY ──────────────────────────────────
    add_styled_heading(doc, "2. Technical Modernization Strategy", level=1, space_before=18)
    
    p4 = doc.add_paragraph("This table details the shift from the current open prototype codebase to the target secure architecture:")
    p4.paragraph_format.space_after = Pt(10)
    
    # Table 1: Current vs Target
    table1 = doc.add_table(rows=6, cols=3)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Feature Area", "Current Behavior", "Target Implementation Plan"]
    col_widths = [Inches(1.8), Inches(2.35), Inches(2.35)]
    
    # Style Header Row
    hdr_cells = table1.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1F4E79")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Segoe UI'
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    rows_data = [
        ("Session Scope", 
         "All questions are visible to any logged-in user with no time bounds.",
         "ExamSessions created with secure UUIDs. Questions are scoped by exam context with active timers. Access requires entering a valid Exam Key."),
        ("Security Gate",
         "Access is open on standard web browsers. Session state resets on page refresh.",
         "Gated behind PyQt6 User-Agent validation, signed session token verification, and mandatory pre-exam security checks (process/display checks)."),
        ("Media Storage",
         "Raw image and audio binary blobs are stored inline inside MongoDB documents.",
         "GridFS integration. Files are stored in separate chunked collections, keeping the main database query speed optimal."),
        ("API Billing & Quota",
         "Gemini AI evaluation triggers automatically upon student submission.",
         "AI evaluations shifted to an administrative trigger model. Evaluations run on-demand to prevent concurrent rate limits and manage cost."),
        ("Biometrics & Proctoring",
         "No biometric identification, active tracking, or data-loss prevention.",
         "Enrolls face templates (Front, Left, Right profiles). Verifies identity before exam start via Gemini. Enforces background camera audits, moving dynamic watermarks, keystroke dynamics, clipboard sanitization, and offline SQLite caching.")
    ]
    
    for row_idx, data in enumerate(rows_data, start=1):
        row_cells = table1.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=120, right=120)
            # Alternating background colors
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], "F2F2F2")
            else:
                set_cell_background(row_cells[col_idx], "FFFFFF")
            p = row_cells[col_idx].paragraphs[0]
            for run in p.runs:
                run.font.name = 'Segoe UI'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(51, 51, 51)
                
    doc.add_page_break()
    
    # ── 3. TARGET ARCHITECTURE & END-TO-END FLOW ───────────────────────────
    add_styled_heading(doc, "3. Target Architecture & End-to-End Flow", level=1, space_before=18)
    
    p5 = doc.add_paragraph("The upgraded system operates through three primary layers: the Desktop Lockdown App (UI Container), the Web Portal (Business Logic & Proctoring API), and the Database (State Persistence).")
    p5.paragraph_format.space_after = Pt(12)
    
    # Steps
    steps = [
        ("Biometric Enrollment", "During administrator student registration, the system initiates a webcam wizard capturing three reference images (Front view, Left profile, Right profile). These templates are stored in GridFS and bound to the user record."),
        ("Exam Initialization", "The administrator creates an exam (assigning questions, duration) and defines a secure alphanumeric Exam Key (access code)."),
        ("Secure Launch", "The user opens the desktop lockdown app in windowed mode and logs in. If they are an administrator, they remain in windowed mode. Students navigate to the instruction gateway."),
        ("Exam Access Gate", "The student selects an exam and inputs the required Exam Key. If valid, the portal unlocks the pre-exam verification phase."),
        ("Security Prechecks", "The PyQt6 client runs automated system checks (ensuring display count is exactly 1, no blacklisted apps are running, and accessibility hooks are available)."),
        ("Identity Verification", "The student takes a live snapshot via the web portal's camera interface. The server transmits this snapshot and the student's enrolled face templates to the Gemini API, which executes biometric structural comparison and returns a Match/Mismatch verdict."),
        ("Token Handshake", "Upon key validation and biometric precheck approvals, the student clicks 'Start Exam'. The PyQt6 client triggers ?action=launch&exam_id=... to initialize the ExamSession UUID inside the database, returning the token to the client app."),
        ("Lockdown Activation", "Once the session token is received, the PyQt6 app immediately intercepts this state, enters frameless fullscreen mode, sets stays-on-top attributes, and activates key suppression hooks."),
        ("Exam Execution", "The page renders only questions assigned to this specific session, initiating a visual countdown timer. The screen displays a moving, transparent watermark of the student's details."),
        ("Proctoring & Heartbeats", "During the active exam, the PyQt6 container monitors system focus, screen count, and processes, forwarding any violations to the API while issuing a silent heartbeat request every 30 seconds. The PyQt6 app runs random background camera audits, clearing the system clipboard, and checking keystroke anomalies."),
        ("Offline Caching", "If internet connection drops, the PyQt6 client autosaves typed responses into an encrypted local SQLite database. Answers are synchronized to MongoDB once the heartbeat connection is restored."),
        ("Exam Completion & Release", "On exam end (timer or manual submit), the client exits fullscreen mode, releases keyboard and monitor hooks, returns the application to a normal windowed container, and gates further submission edits."),
        ("Administrative AI Grading", "The admin reviews proctoring logs and runs AI grading via the Gemini API on-demand.")
    ]
    
    for idx, (title, desc) in enumerate(steps, start=1):
        step_p = doc.add_paragraph()
        step_p.paragraph_format.left_indent = Inches(0.25)
        step_p.paragraph_format.space_after = Pt(4)
        run_idx = step_p.add_run(f"{idx}. {title}: ")
        run_idx.bold = True
        run_idx.font.name = 'Segoe UI'
        run_idx.font.size = Pt(10.5)
        run_idx.font.color.rgb = RGBColor(31, 78, 121)
        
        run_desc = step_p.add_run(desc)
        run_desc.font.name = 'Segoe UI'
        run_desc.font.size = Pt(10.5)
        run_desc.font.color.rgb = RGBColor(64, 64, 64)
        
    doc.add_page_break()
    
    # ── 4. DETAILED COMPONENT SPECIFICATIONS ─────────────────────────────────
    add_styled_heading(doc, "4. Component Specifications", level=1, space_before=18)
    
    # 4.1 Database Layer
    add_styled_heading(doc, "4.1 Database Enhancements (db.py)", level=2)
    p_db = doc.add_paragraph("To support proctoring, key gating, and biometrics, collections are updated:")
    p_db.paragraph_format.space_after = Pt(6)
    
    collections_data = [
        ("users", "Updated to store face profile references: face_templates (array of GridFS ObjectIDs matching enrolled Front, Left, and Right profiles)."),
        ("exams", "Updated to store configuration data: question_ids, title, duration_minutes, created_by, is_active, and access_key (hashed alphanumeric passcode)."),
        ("exam_sessions", "Tracks student sessions: token (UUID), started_at, expires_at, status, violations (array of logged alerts), last_heartbeat.")
    ]
    for col_name, col_desc in collections_data:
        cp = doc.add_paragraph(style='List Bullet')
        cp.paragraph_format.space_after = Pt(3)
        rc_bold = cp.add_run(f"{col_name}: ")
        rc_bold.bold = True
        rc_bold.font.name = 'Segoe UI'
        cp.add_run(col_desc).font.name = 'Segoe UI'
        
    p_db_util = doc.add_paragraph("GridFS is initialized using PyMongo's gridfs module. Binary file upload functions are refactored to write to GridFS first, saving the returned file ObjectIDs inside the submission document instead of storing inline BSON bytes. Database connection logic (get_db) is updated with ping commands to automatically verify host server connectivity and reconnect dynamically on failure.")
    p_db_util.paragraph_format.space_before = Pt(8)
    p_db_util.paragraph_format.space_after = Pt(12)
    
    # 4.2 Portal Routing
    add_styled_heading(doc, "4.2 Web Portal Routing (app.py)", level=2)
    p_route = doc.add_paragraph("Custom CSS is moved from the main application script into an independent styles.py module, exposing a clean stylesheet function to minimize code bloat. An API parameter router is added to the top of main() to parse query parameters before page layouts render:")
    p_route.paragraph_format.space_after = Pt(6)
    
    routes = [
        ("?action=launch&u=<username>&exam_id=<exam_id>", "Initializes the database session for the specified exam after validation checks, outputting the token JSON payload."),
        ("?action=heartbeat&token=<token>", "Updates the session heartbeat timestamp. Used by the desktop app keep-alive timer."),
        ("?action=violation&token=<token>&type=<type>", "Logs proctoring violation alerts directly into the session document."),
        ("?action=submit&token=<token>", "Concludes the active exam session and locks the user interface."),
        ("?page=upload&token=<token>", "A public page bypassing standard auth gates, allowing mobile companion camera uploads.")
    ]
    for r_name, r_desc in routes:
        rp = doc.add_paragraph(style='List Bullet')
        rp.paragraph_format.space_after = Pt(3)
        rr_bold = rp.add_run(f"{r_name}: ")
        rr_bold.bold = True
        rr_bold.font.name = 'Segoe UI'
        rp.add_run(r_desc).font.name = 'Segoe UI'
        
    # 4.3 Student Dashboard
    add_styled_heading(doc, "4.3 Student Dashboard Upgrades (student_pages.py)", level=2)
    p_student = doc.add_paragraph("The student experience is updated to enforce exam integrity:")
    p_student.paragraph_format.space_after = Pt(6)
    
    student_features = [
        ("Access Verification Gate", "Checks that the incoming connection uses the PyQt6 User-Agent and possesses a valid token query parameter. Fails if standard browsers are used."),
        ("Exam Key Validation Screen", "Prompts the student to select their target exam and enter the Exam Key. Access is denied if the key does not match."),
        ("Pre-Exam Biometric Verification Screen", "Initiates the webcam via st.camera_input() and requires the student to capture a live face photo. This snapshot is compared against the enrolled profiles. Exam launch is blocked unless Gemini API issues a verified match confirmation."),
        ("Dynamic Watermarking Overlay", "Injects moving, semi-transparent text blocks containing the student's username, current local IP, and a rolling timestamp. Acts as a visual barrier preventing screen photography leaks."),
        ("Keyboard Injection Monitor", "A JavaScript hook logs the keystroke intervals (dwell and flight timings). If characters are entered at speeds exceeding human typing boundaries, the system logs a keystroke_injection violation.")
    ]
    for sf_name, sf_desc in student_features:
        sfp = doc.add_paragraph(style='List Bullet')
        sfp.paragraph_format.space_after = Pt(3)
        rsf_bold = sfp.add_run(f"{sf_name}: ")
        rsf_bold.bold = True
        rsf_bold.font.name = 'Segoe UI'
        sfp.add_run(sf_desc).font.name = 'Segoe UI'
        
    # 4.4 Admin Panel
    add_styled_heading(doc, "4.4 Proctoring Dashboard & Grading (admin_pages.py)", level=2)
    p_admin = doc.add_paragraph("The admin portal is equipped with oversight capabilities:")
    p_admin.paragraph_format.space_after = Pt(6)
    
    admin_features = [
        ("Biometric Enrollment Wizard", "Integrated into the Add Student form. Features a multi-step webcam capture sequence to register and store 3 distinct reference angles (Front, Left, Right) inside the user's database document."),
        ("Live Proctoring Dashboard", "A dedicated tab to track student status, display real-time violation logs (focus loss, process violations, double screen, camera audit flags, keyboard injection events), heartbeat timestamps, and a manual session termination button."),
        ("AI Grading Management", "Shifts AI analysis from submission time to administrative control. Each entry features a 'Trigger Gemini Evaluation' button to run grading on-demand.")
    ]
    for af_name, af_desc in admin_features:
        afp = doc.add_paragraph(style='List Bullet')
        afp.paragraph_format.space_after = Pt(3)
        raf_bold = afp.add_run(f"{af_name}: ")
        raf_bold.bold = True
        raf_bold.font.name = 'Segoe UI'
        afp.add_run(af_desc).font.name = 'Segoe UI'
        
    # 4.5 Desktop App
    add_styled_heading(doc, "4.5 PyQt6 Desktop App Specifications (lockdown_browser.py)", level=2)
    p_desktop = doc.add_paragraph("Built in PyQt6, this application acts as a secure container for the exam:")
    p_desktop.paragraph_format.space_after = Pt(6)
    
    desktop_features = [
        ("Dynamic Window States", "The application starts as a standard windowed window. Fullscreen proctoring and keyboard hooks are only engaged when the page URL indicates an exam has started (e.g. ?page=exam). This allows administrators to manage exams, and students to read instructions, without entering lockdown prematurely."),
        ("Low-Level Keyboard Hook (Windows)", "Active only in lockdown mode. Uses a ctypes WH_KEYBOARD_LL hook to intercept and block system shortcuts (Windows key, Alt+Tab, Alt+Esc, Ctrl+Esc)."),
        ("Cocoa Presentation Controls (macOS)", "Active only in lockdown mode. Disables dock hiding, process switching, and the force-quit shortcut menu via native NSApplication overrides."),
        ("Continuous Proctoring Monitors", "Checks display count every 2 seconds. The Focus Monitor detects focus loss. The Process Monitor checks background processes against a whitelist every 10 seconds. Violations are instantly sent to the server proctoring API when lockdown is active."),
        ("Continuous Camera Audits", "A background thread uses QCamera to capture webcam frames at random intervals (3–5 minutes). The frames are uploaded to the proctoring API for Gemini evaluation (detecting face mismatch, multiple people, or smartphones/books)."),
        ("Clipboard Sanitizer", "Wipes the operating system clipboard buffer upon entering lockdown and blocks copy/paste shortcuts."),
        ("Local SQLite Cache", "Maintains an encrypted SQLite file. Typed responses are cached locally every 10 seconds. In the event of network connection failure, the app continues to record answers offline, syncing with MongoDB once the heartbeat recovers."),
        ("Crash-to-Resume Recovery", "Saves the session token in an encrypted state file. If the app restarts unexpectedly, the student can click 'Resume Exam', complete a new face validation check, and continue with their remaining time.")
    ]
    for df_name, df_desc in desktop_features:
        dfp = doc.add_paragraph(style='List Bullet')
        dfp.paragraph_format.space_after = Pt(3)
        rdf_bold = dfp.add_run(f"{df_name}: ")
        rdf_bold.bold = True
        rdf_bold.font.name = 'Segoe UI'
        dfp.add_run(df_desc).font.name = 'Segoe UI'

    # 4.6 Gemini Biometric Verification
    add_styled_heading(doc, "4.6 Gemini Biometric & Proctoring Evaluation Engine (gemini_eval.py)", level=2)
    p_gemini_bio = doc.add_paragraph("A new AI module is introduced to handle user identity verification. Instead of importing complex computer vision local libraries that could break across target machines, the system leverages Gemini's multi-modal capabilities to perform verification:")
    p_gemini_bio.paragraph_format.space_after = Pt(6)

    bio_steps = [
        ("Face Matching Payload", "The portal reads the three enrolled face templates from GridFS and packs them alongside the live pre-exam verification snapshot as multi-modal image blocks. Gemini is prompted to compare facial details and return a Match/Mismatch verdict."),
        ("Continuous Camera Audits", "Gemini receives random background proctoring snapshots. It evaluates for face matching, checks for additional individuals in the background, and flags unauthorized objects like cell phones, reference books, or tablets.")
    ]
    for bs_name, bs_desc in bio_steps:
        bsp = doc.add_paragraph(style='List Bullet')
        bsp.paragraph_format.space_after = Pt(3)
        rbs_bold = bsp.add_run(f"{bs_name}: ")
        rbs_bold.bold = True
        rbs_bold.font.name = 'Segoe UI'
        bsp.add_run(bs_desc).font.name = 'Segoe UI'
        
    # 4.7 API Cost & Quota Optimization
    add_styled_heading(doc, "4.7 API Cost & Quota Optimization (Computational Offloading)", level=2)
    p_opt = doc.add_paragraph("Running constant multi-modal analysis via Gemini can lead to high request latency and rate limit exhausts. The plan incorporates three mitigation strategies:")
    p_opt.paragraph_format.space_after = Pt(6)

    opts_steps = [
        ("Local Frame-Differencing", "The PyQt6 client runs a lightweight pixel comparison using the Pillow library on successive frames. If pixel variance between frames is under 5% (indicating the student is sitting still), the upload is skipped, reducing API traffic by 80%."),
        ("Resolution Downscaling", "Webcam snapshots are compressed to a resolution of 320x240 and saved in JPEG format at 60% quality prior to upload. This decreases the input token footprint to a negligible size, maintaining minimum costs."),
        ("Randomized Audit Throttling", "The client runs proctoring checks continuously but only selects 3 to 5 random snapshots per hour to send to the Gemini API for biometric verification, keeping a strong proctoring deterrent while staying well within the free-tier API quotas.")
    ]
    for os_name, os_desc in opts_steps:
        osp = doc.add_paragraph(style='List Bullet')
        osp.paragraph_format.space_after = Pt(3)
        ros_bold = osp.add_run(f"{os_name}: ")
        ros_bold.bold = True
        ros_bold.font.name = 'Segoe UI'
        osp.add_run(os_desc).font.name = 'Segoe UI'
        
    doc.add_page_break()
    
    # ── 5. TIMELINE & MILESTONES ─────────────────────────────────────────────
    add_styled_heading(doc, "5. Timeline & Project Milestones", level=1, space_before=18)
    
    p_timeline = doc.add_paragraph("The implementation timeline is structured into six phases spanning ~45 working days (~9 calendar weeks). This estimate is tailored for a 3rd-year BTech Computer Science & Engineering intern, incorporating a 25% project buffer to support thorough testing, debugging, and review cycles.")
    p_timeline.paragraph_format.space_after = Pt(12)
    
    # Timeline Table
    table2 = doc.add_table(rows=7, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers_t = ["Phase", "Focus Area", "Raw Estimate", "Buffered Estimate"]
    col_widths_t = [Inches(1.2), Inches(2.8), Inches(1.2), Inches(1.3)]
    
    hdr_cells_t = table2.rows[0].cells
    for i, title in enumerate(headers_t):
        hdr_cells_t[i].text = title
        set_cell_background(hdr_cells_t[i], "1F4E79")
        set_cell_margins(hdr_cells_t[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells_t[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Segoe UI'
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    t_data = [
        ("Phase 1", "Foundation & Data Models (reconnection logic, UTC standardization, GridFS, new collections, CSS refactoring)", "8 Days", "10 Days"),
        ("Phase 2", "Portal Exam & Admin Dashboards (exam setup with keys, face enrollment wizard, proctoring monitor view, routing, token verification)", "10 Days", "12 Days"),
        ("Phase 3", "Security Prechecks & Biometric Match Gate (exam key gates, precheck runs, st.camera biometric capture, Gemini face verification API)", "8 Days", "10 Days"),
        ("Phase 4", "Desktop Lockdown App Core (PyQt6 window, keyboard hooks, multi-screen blocks, process scans, heartbeat pings, URL change triggers)", "10 Days", "12 Days"),
        ("Phase 5", "Advanced Proctoring & Resilience (PyQt6 background camera capture, dynamic watermark overlay, clipboard sanitizer, offline SQLite cache & resume flow, local differencing filters)", "8 Days", "10 Days"),
        ("Phase 6", "Testing, Packaging & Staging (QA, PyInstaller builds for Windows/macOS, README & security settings configuration)", "6 Days", "8 Days")
    ]
    
    for row_idx, data in enumerate(t_data, start=1):
        row_cells = table2.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=120, right=120)
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], "F2F2F2")
            else:
                set_cell_background(row_cells[col_idx], "FFFFFF")
            p = row_cells[col_idx].paragraphs[0]
            for run in p.runs:
                run.font.name = 'Segoe UI'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(51, 51, 51)
                
    doc.add_page_break()
    
    # ── 6. SYSTEM BOUNDARIES & LIMITATIONS ───────────────────────────────────
    add_styled_heading(doc, "6. Operating System Boundaries & Limitations", level=1, space_before=18)
    
    p_lim = doc.add_paragraph("This section highlights standard operating system boundaries and deployment considerations:")
    p_lim.paragraph_format.space_after = Pt(8)
    
    lims = [
        ("Ctrl+Alt+Del (Windows)", "Ctrl+Alt+Del is a hardware interrupt handled directly by the Windows kernel security architecture. Blocking this key sequence at the application level is not supported by standard APIs and requires kernel-level driver development or an active Windows Service. This is out of scope for the application layer of this build and is documented as a native OS boundary."),
        ("Accessibility Permissions (macOS)", "To monitor keyboard events and override system window controls, macOS requires the user to grant explicit Accessibility Permissions in System Settings on first launch. The application will prompt the user to enable this setting but cannot toggle it automatically."),
        ("QR Code Network Requirements", "For local-only development and testing environments, the QR companion upload page requires the student's phone and the server host to reside on the same Wi-Fi/local network. For staging and production setups, exposing the server via a public domain or hosting gateway solves this network constraint.")
    ]
    
    for l_title, l_desc in lims:
        lp = doc.add_paragraph(style='List Bullet')
        lp.paragraph_format.space_after = Pt(6)
        rl_bold = lp.add_run(f"{l_title}: ")
        rl_bold.bold = True
        rl_bold.font.name = 'Segoe UI'
        lp.add_run(l_desc).font.name = 'Segoe UI'
        
    try:
        doc.save('EduPortal_AI_Implementation_Plan.docx')
        print("Regenerated EduPortal_AI_Implementation_Plan.docx successfully.")
    except PermissionError:
        doc.save('EduPortal_AI_Implementation_Plan_v2.docx')
        print("Could not overwrite EduPortal_AI_Implementation_Plan.docx (file locked). Saved as EduPortal_AI_Implementation_Plan_v2.docx successfully.")

if __name__ == '__main__':
    create_document()
